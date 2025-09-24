import json
from django.contrib.auth import authenticate, login, logout
from django.http import JsonResponse
from django.middleware.csrf import get_token
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.http import require_POST
from django.contrib.sessions.models import Session
from rest_framework.renderers import JSONRenderer
from django.http import FileResponse

from .serializers import PatientSerializer
from .serializers import VisitSerializer
from .models import Patient
from .models import Visit

# Декоратор для выдачи ошибки если пользователь неавторизован
def json_login_required(view_func):
    def wrapped_view(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return JsonResponse({'error': 'Вы не авторизованы'}, status=401)
        return view_func(request, *args, **kwargs)
    return wrapped_view

def get_csrf(request):
    response = JsonResponse({'detail': 'CSRF cookie set'})
    s = get_token(request)
    response['X-CSRFToken'] = s
    return response

@require_POST
def login_view(request):
    # Получаем авторизационные данные
    data = json.loads(request.body)
    username = data.get('username')
    password = data.get('password')

    # Валидация
    if username is None or password is None:
        return JsonResponse({'detail': 'Пожалуйста предоставьте логин и пароль'}, status=400)

    # Аутентификация пользоваля
    user = authenticate(username=username, password=password)
    
    if user is None:
        return JsonResponse({'detail': 'Неверный логин или пароль'}, status=400)

    # Создаётся сессия. session_id отправляется в куки
    login(request, user)
    return JsonResponse({'detail': 'Успешная авторизация'})

  
# Сессия удаляется из БД и session_id на клиенте более недействителен
@json_login_required
def logout_view(request):
    logout(request)
    return JsonResponse({'detail': 'Вы успешно вышли'})

  
# Узнать авторизован ли пользователь и получить его данные
@ensure_csrf_cookie # <- Принудительная отправка CSRF cookie
def session_view(request):
    if not request.user.is_authenticated:
        return JsonResponse({'isAuthenticated': False})

    return JsonResponse({'isAuthenticated': True, 'username': request.user.username, 'user_id': request.user.id})

  
# Получение информации о пользователе
@json_login_required
def user_info(request):
    return JsonResponse({'username': request.user.username, 'firstname': request.user.first_name, 'lastname': request.user.last_name})

@json_login_required
def get_patients(request):
    patients_from_db = Patient.objects.all()
    serialized = PatientSerializer(patients_from_db, many=True) 
    return JsonResponse(serialized.data,safe=False)

@json_login_required
def test(request,id):
    my_object = Patient.objects.filter(pk=id).first()
    return JsonResponse(PatientSerializer(my_object).data)

@json_login_required
def get_analysis(request):
    tapID = request.GET.get("tapID")
    visits = Visit.objects.filter(tapID=tapID)
    serialized = VisitSerializer(visits, many=True) 
    return JsonResponse(serialized.data, safe=False)




# Удаление всех сессий из БД
# Вы можете переделать так, чтобы отзывать сессию у определённого пользователя
@json_login_required
def kill_all_sessions(request):
    sessions = Session.objects.all()
    sessions.delete()

    return JsonResponse({'detail': 'Сессии успешно завершены'})

import pandas as pd
import pickle
def diagnose_predict(request):
    codes = [
        "D50.9 Железодефицитная анемия неуточненная",
        "D51.9 Витамин-B12-дефицитная анемия неуточненная",
        "D52.9 Фолиеводефицитная анемия неуточненная",
        "D53.9 Анемия, связанная с питанием, неуточненная", 
        "D56.9 Талассемия неуточненная", 
        "D57.8 Другие серповидно-клеточные нарушения",
        "D58.9 Наследственная гемолитическая анемия неуточненная",
        "D59.9 Приобретенная гемолитическая анемия неуточненная", 
        "D60.9 Приобретенная чистая красно-клеточная аплазия неуточненная",
        "D61.9 Апластическая анемия неуточненная",
        "D62.0 Острая постгеморрагическая анемия",
        "D63.8 Анемия при других хронических болезнях, классифицированных в других рубриках",
        "D64.4 Врожденная дизэритропоэтическая анемия"
    ]

    pol = request.GET.get('pol')
    ves = request.GET.get('ves')
    travma = request.GET.get('travma')
    
    visits = Visit.objects.filter(tapID=request.GET.get("tap"))
    analyses = {}

    for visit in visits:
        analyses[visit.investigationName] = visit.investigationResult

    cols = ["pol", "ves", "travma", "onko", "infec", "uzi", "nasled",
        "HGB", "erit", "leik", "PLT", "gematok", "MCH", "MCHC", "MCV",
        "pokazatel", "Fe", "OZSS", "Ferrit", "B12", "billirubin", "belok",
        "folievay", "albumin", "mielogramma", "Kumbs"
    ]

    gender = 0 if request.GET.get('pol')=="жен" else 1
    weight = float(request.GET.get('ves'))/100
    array = pd.DataFrame([
        [gender, #Пол
        weight, #Вес
        request.GET.get('travma'), #травма 
        request.GET.get('onko'), #онко +
        request.GET.get('infec'), #инфекция +  
        request.GET.get('uzi'), #узи +
        request.GET.get('nasled'), #наследст +
        analyses["HGB"], #HGB Исследование уровня общего гемоглобина в крови
        analyses["erit"], # erit Исследование уровня эритроцитов в крови
        analyses["leik"], # leik Исследование уровня лейкоцитов в крови
        analyses["PLT"], #PLT Исследование уровня тромбоцитов в крови
        analyses["gematok"], #gematok Оценка гематокрита
        analyses["MCH"], #MCH Определение среднего содержания гемоглобина в эритроцитах (MCH)	
        analyses["MCHC"],	#MCHC Средняя концентрация гемоглобина в эритроцитах (MCHC)
        analyses["MCV"], #MCV Средний объем эритроцитов (MCV)
        analyses["pokazatel"], #Цветовой показатель
        analyses["Fe"], #Fe Исследование уровня железа в сыворотке крови
        analyses["OZSS"], #OZSS Исследование железосвязывающей способности сыворотки (ОЖСС)
        analyses["Ferrit"], #Ferrit Исследование уровня ферритина в крови
        analyses["B12"], #B12 Определение уровня витамина B12 (цианокобаламин) в крови
        analyses["billirubin"], #billirubin Исследование уровня общего билирубина в крови
        analyses["belok"], #belok Исследование уровня общего белка в крови
        analyses["folievay"], #folievay Исследование уровня фолиевой кислоты в сыворотке крови
        analyses["albumin"], #albumin Исследование уровня альбумина в крови
        analyses["mielogramma"], #mielogramma (миелограмма)
        analyses["Kumbs"] #Прямой антиглобулиновый тест (прямая проба Кумбса) 
        ], 
    ],columns=cols)
    

    with open("ml/model.pkl", "rb") as f:
        model = loaded_model = pickle.load(f)
        res = model.predict(array)
        result = codes[res[0]-1]

    return JsonResponse({'result':result})


def get_document(request):
    from docx import Document
    from datetime import datetime;

    now = datetime.now()

    date = str(now.day)+"."+str(now.month)+"."+str(now.year)

    doc = Document()
    doc.add_heading('Отчёт о посещении', level=1)
    doc.add_paragraph('Дата посещения: '+date)
    
    doc.add_paragraph('ФИО врача: '+request.GET.get('doctor'))

    doc.add_paragraph('Основной диагноз: '+request.GET.get('base_diag'))
    doc.add_paragraph('Диагноз системы: '+request.GET.get('system_diag'))

    # doc.add_paragraph('Диагноз: '+date)
    names2 = [
        'Пол',
        'Вес',
        'Наличие кровотечения',
        'Наличие на момент обследования восполительных заболеваний, онкологий',
        'Инфекции, переливание крови, отравление, интоксикация',  
        'Результат ультразвукового исследования'
        'Исследование уровня общего гемоглобина в крови',
        'Исследование уровня эритроцитов в крови',
        'Исследование уровня лейкоцитов в крови',
        'Исследование уровня тромбоцитов в крови',
        'Оценка гематокрита',
        'Определение среднего содержания гемоглобина в эритроцитах MCH',
        'Средняя концентрация гемоглобина в (MCHC)',
        'Средний объем эритроцитов MCV',
        'Цветовой  показатель',
        'Исследование уровня железа в сыворотки крови',
        'Исследование железосвязывающей способности сыворотки (ОЖСС)',
        'Исследование уровня ферритина в крови',
        'Определение уровня витамина B12 (цианокобаламин) в крови',
        'Исследование уровня общего билирубина в крови',
        'Исследование уровня общего белка в крови',
        'Исследование уровня фолиевой кислоты в сыворотке крови',
        'Исследование уровня альбумина в крови',
        'Цитологическое исследование мазка костного мозга (миелограмма)',
        'Прямой антиглобулиновый тест (прямая проба Кумбса)'
    ]

    normals = [
        '-',
        '-',
        '-',
        '-',
        '-',
        '-',
        '120-140 (ж), '+'135-160 (м)',
        '3,9-4,7 (ж), '+'4-5 (м)',
        '4-9',
        '150-400',
        '36-42 (ж),  '+'40-41 (м)',
        '24-34',
        '300-380',
        '75-95',
        'Вычисляется по формуле MCH*0,03 (формулу выводить на экран не надо)',
        '6,6 – 26',
        '11 – 28',
        '45,3-77,1',
        '10-120 (М)  ' +'20-250 (Ж)',
        '197-771',
        '2-21',
        '65-85',
        '3 - 17',
        '35-53',
        '0 - норма  '+'1-угнетение более 2-х ростков  '+'  0,5 - изолированное угнетение красного ростка',
        '0-отриц'
    ]

    units = [
        '-',
        '-',
        '-',
        '-',
        '-',
        '-',
        '-',
        '*10^12 на л (м)',
        '*10^9/л',
        '*10^9/л',
        '-',
        '-',
        '-',
        '-',
        '-',
        'мкмоль/л',
        'мкмоль/л',
        'мкг/л',
        'пг/мл',
        'мкмоль/л',
        'г/л',
        'нг/мл',
        'г/л',
        '-',
        '-'
    ]

    


    pol = request.GET.get('pol')
    ves = request.GET.get('ves')
    travma = request.GET.get('travma')

    visits = Visit.objects.filter(tapID=request.GET.get("tap"))
    analyses = {}
    for visit in visits:
        analyses[visit.investigationName] = visit.investigationResult


    patient_values = [
        pol, #Пол
        ves, #Вес
        "Да" if request.GET.get('travma')=='1' else "Нет", #травма 
        "Да" if request.GET.get('onko') else "Нет", #онко +
        "Да" if request.GET.get('infec') else "Нет", #инфекция +  
        "Отклонения есть" if request.GET.get('uzi') else "Отклонений нет", #узи +
        # "Да" if request.GET.get('nasled') else "Нет", #наследст +
        analyses["HGB"], #HGB Исследование уровня общего гемоглобина в крови
        analyses["erit"], # erit Исследование уровня эритроцитов в крови
        analyses["leik"], # leik Исследование уровня лейкоцитов в крови
        analyses["PLT"], #PLT Исследование уровня тромбоцитов в крови
        analyses["gematok"], #gematok Оценка гематокрита
        analyses["MCH"], #MCH Определение среднего содержания гемоглобина в эритроцитах (MCH)	
        analyses["MCHC"],	#MCHC Средняя концентрация гемоглобина в эритроцитах (MCHC)
        analyses["MCV"], #MCV Средний объем эритроцитов (MCV)
        analyses["pokazatel"], #Цветовой показатель
        analyses["Fe"], #Fe Исследование уровня железа в сыворотке крови
        analyses["OZSS"], #OZSS Исследование железосвязывающей способности сыворотки (ОЖСС)
        analyses["Ferrit"], #Ferrit Исследование уровня ферритина в крови
        analyses["B12"], #B12 Определение уровня витамина B12 (цианокобаламин) в крови
        analyses["billirubin"], #billirubin Исследование уровня общего билирубина в крови
        analyses["belok"], #belok Исследование уровня общего белка в крови
        analyses["folievay"], #folievay Исследование уровня фолиевой кислоты в сыворотке крови
        analyses["albumin"], #albumin Исследование уровня альбумина в крови
        analyses["mielogramma"], #mielogramma (миелограмма)
        analyses["Kumbs"] #Прямой антиглобулиновый тест (прямая проба Кумбса) 
    ] 


    table = doc.add_table(rows=len(names2)+1, cols=4)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Данные'
    hdr_cells[1].text = 'Норма'
    hdr_cells[2].text = 'Ед. измерения'
    hdr_cells[3].text = 'Пациент'

    print(len(names2)+1)
    for i in range(1, len(names2)+1):
        row_cells = table.rows[i].cells
        row_cells[0].text = names2[i-1]
        row_cells[1].text = normals[i-1]
        row_cells[2].text = units[i-1]
        if patient_values[i-1] is None:
            row_cells[3].text = ""
        else:
            row_cells[3].text = patient_values[i-1]
        

    doc.save('backend/documents/order.docx')
    file_path = "backend/documents/order.docx"

  
    response = FileResponse(open(file_path, 'rb'), as_attachment=True, filename='PatientResult.docx')
    # response['Content-Disposition'] = 'attachment; filename="name.docx"'
    return response 

    